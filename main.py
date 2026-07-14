# -*- coding: utf-8 -*-
"""
ANÁLISIS BASAL DE PRESIÓN ARTERIAL - V06 DASHBOARD CLÍNICO ESH PRO

Mejoras principales:
- Pregunta el número de días para la línea base.
- Lista las fechas/días disponibles para seleccionar el inicio.
- Aplica orientación interpretativa basada en criterios domiciliarios ESH 2023.
- Genera dashboard profesional por sujeto.
- Separa la presión arterial y la frecuencia cardiaca en gráficas distintas.
- Sustituye la tabla horizontal excesiva por una tabla vertical:
  Sigla | Nombre completo | Valor | Interpretación | Recomendaciones.

Estructura esperada:
    /Datos
        Pepe_Pino.xlsx
    /Resultados
        Informe_Basal_Presion_Arterial_V06_ESH.docx
        Resultados_Basal_Presion_Arterial_V06_ESH.xlsx
        /Graficos

Formato recomendado del Excel:
    Día | Fecha | Hora | Medición | PAS | PAD | Frecuencia cardiaca | Variabilidad FC

Instalar dependencias:
    py -m pip install pandas openpyxl python-docx matplotlib numpy
"""

from pathlib import Path
from datetime import datetime
import re
import sys
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


CARPETA_DATOS = Path("data/input")
CARPETA_RESULTADOS = Path("results")
CARPETA_GRAFICOS = CARPETA_RESULTADOS / "figures"
CARPETA_REPORTES = CARPETA_RESULTADOS / "reports"
CARPETA_CONVERTIDOS = Path("data/output/converted")

NOMBRE_WORD = "Informe_Basal_Presion_Arterial_V06_Dashboard_Clinico_ESH_PRO.docx"
NOMBRE_EXCEL = "Resultados_Basal_Presion_Arterial_V06_Dashboard_Clinico_ESH_PRO.xlsx"

# Orientación domiciliaria ESH 2023: HTA domiciliaria si PAS media >=135 y/o PAD media >=85 mmHg.
UMBRAL_HOME_PAS = 135
UMBRAL_HOME_PAD = 85
UMBRAL_URGENTE_PAS = 180
UMBRAL_URGENTE_PAD = 120


def normalizar_texto(x):
    if pd.isna(x):
        return ""
    x = str(x).strip().lower()
    x = x.replace("\n", " ")
    for a, b in [("á", "a"), ("é", "e"), ("í", "i"), ("ó", "o"), ("ú", "u"), ("ü", "u")]:
        x = x.replace(a, b)
    x = re.sub(r"\s+", " ", x)
    return x


def detectar_columna(columnas, patrones):
    for col in columnas:
        col_limpia = normalizar_texto(col)
        for patron in patrones:
            if normalizar_texto(patron) in col_limpia:
                return col
    return None


def extraer_nombre_sujeto(ruta):
    return ruta.stem.replace("_", " ").replace("-", " ").strip()


def limpiar_numero(serie):
    return pd.to_numeric(
        serie.astype(str)
        .str.replace(",", ".", regex=False)
        .str.replace(" ", "", regex=False),
        errors="coerce"
    )


def leer_excel(ruta_excel):
    df = pd.read_excel(ruta_excel)
    df = df.dropna(how="all").dropna(axis=1, how="all").copy()

    normalizadas = {col: normalizar_texto(col) for col in df.columns}

    def exacta(*nombres):
        objetivos = {normalizar_texto(n) for n in nombres}
        for col, nombre in normalizadas.items():
            if nombre in objetivos:
                return col
        return None

    col_fecha_hora = exacta("fecha/hora", "fecha hora", "datetime", "timestamp")
    col_dia = exacta("día", "dia", "day")
    col_fecha = exacta("fecha", "date")
    col_hora = exacta("hora", "time")
    col_medicion = detectar_columna(df.columns, ["medicion", "medición", "toma"])

    col_pas = exacta("sys(mmhg)", "sys", "pas") or detectar_columna(
        df.columns, ["sys(mmhg)", "systolic", "sistolica", "sistólica", "pas"]
    )
    col_pad = exacta("dia(mmhg)", "pad") or detectar_columna(
        df.columns, ["dia(mmhg)", "diastolic", "diastolica", "diastólica", "pad"]
    )
    col_fc = exacta("pr(bpm)", "pr", "fc", "hr") or detectar_columna(
        df.columns,
        ["pr(bpm)", "frecuencia cardiaca", "frecuencia cardíaca", "pulso", "heart rate", "fc", "hr"],
    )
    col_hrv = detectar_columna(df.columns, ["variabilidad", "hrv", "rmssd", "lnrmssd"])
    col_sujeto = detectar_columna(df.columns, ["nombre o id", "nombre", "sujeto", "subject", "id"])

    faltan = []
    if col_pas is None:
        faltan.append("PAS/SYS")
    if col_pad is None:
        faltan.append("PAD/DIA")
    if col_fc is None:
        faltan.append("FC/PR")
    if faltan:
        raise ValueError("No se detectaron columnas obligatorias: " + ", ".join(faltan))

    out = pd.DataFrame(index=df.index)

    if col_sujeto is not None:
        sujetos = df[col_sujeto].dropna().astype(str).str.strip()
        sujeto = sujetos.iloc[0] if not sujetos.empty else extraer_nombre_sujeto(ruta_excel)
    else:
        sujeto = extraer_nombre_sujeto(ruta_excel)

    out["Sujeto"] = sujeto
    out["Archivo"] = ruta_excel.name

    if col_fecha_hora is not None:
        fecha_hora = pd.to_datetime(df[col_fecha_hora], errors="coerce", dayfirst=True)
        out["Fecha_dt"] = fecha_hora
        out["Fecha"] = fecha_hora.dt.strftime("%d/%m/%Y")
        out["Hora"] = fecha_hora.dt.strftime("%H:%M:%S")
    else:
        if col_fecha is not None:
            out["Fecha"] = df[col_fecha]
            out["Fecha_dt"] = pd.to_datetime(df[col_fecha], errors="coerce", dayfirst=True)
        else:
            out["Fecha"] = ""
            out["Fecha_dt"] = pd.NaT
        out["Hora"] = df[col_hora] if col_hora is not None else ""

    if col_dia is not None:
        out["Dia"] = limpiar_numero(df[col_dia])
    elif out["Fecha_dt"].notna().any():
        fechas = out["Fecha_dt"].dt.normalize()
        fechas_unicas = sorted(fechas.dropna().unique())
        mapa = {fecha: i for i, fecha in enumerate(fechas_unicas, start=1)}
        out["Dia"] = fechas.map(mapa)
    else:
        out["Dia"] = (np.arange(len(df)) // 2) + 1

    if col_medicion is not None:
        out["Medicion"] = limpiar_numero(df[col_medicion])
    else:
        out["Medicion"] = out.groupby("Dia").cumcount() + 1

    out["PAS"] = limpiar_numero(df[col_pas])
    out["PAD"] = limpiar_numero(df[col_pad])
    out["FC"] = limpiar_numero(df[col_fc])
    out["HRV"] = limpiar_numero(df[col_hrv]) if col_hrv is not None else np.nan

    out.loc[(out["PAS"] < 70) | (out["PAS"] > 260), "PAS"] = np.nan
    out.loc[(out["PAD"] < 40) | (out["PAD"] > 160), "PAD"] = np.nan
    out.loc[(out["FC"] < 35) | (out["FC"] > 190), "FC"] = np.nan

    out = out.dropna(subset=["Dia", "PAS", "PAD", "FC"], how="any").copy()
    out["Dia"] = out["Dia"].astype(int)
    return out

def calcular_medias_diarias(df):
    medias = (
        df.groupby(["Sujeto", "Archivo", "Dia"], as_index=False)
        .agg(
            Fecha=("Fecha", "first"),
            Fecha_dt=("Fecha_dt", "first"),
            N_mediciones=("Medicion", "count"),
            PAS_MEDIA=("PAS", "mean"),
            PAD_MEDIA=("PAD", "mean"),
            FC_MEDIA=("FC", "mean"),
            HRV_MEDIA=("HRV", "mean"),
            PAS_1=("PAS", "first"),
            PAS_2=("PAS", "last"),
            PAD_1=("PAD", "first"),
            PAD_2=("PAD", "last"),
            FC_1=("FC", "first"),
            FC_2=("FC", "last"),
        )
    )
    medias["DIF_PAS_1_2"] = medias["PAS_1"] - medias["PAS_2"]
    medias["DIF_PAD_1_2"] = medias["PAD_1"] - medias["PAD_2"]
    medias["DIF_FC_1_2"] = medias["FC_1"] - medias["FC_2"]
    medias["PRESION_PULSO"] = medias["PAS_MEDIA"] - medias["PAD_MEDIA"]
    medias["PAM"] = medias["PAD_MEDIA"] + ((medias["PAS_MEDIA"] - medias["PAD_MEDIA"]) / 3)
    return medias.sort_values(["Sujeto", "Dia"]).copy()


def pedir_entero(mensaje, minimo=1, maximo=None, defecto=None):
    while True:
        entrada = input(mensaje).strip()
        if entrada == "" and defecto is not None:
            return defecto
        try:
            valor = int(entrada)
            if valor < minimo:
                print(f"Introduzca un valor >= {minimo}.")
                continue
            if maximo is not None and valor > maximo:
                print(f"Introduzca un valor <= {maximo}.")
                continue
            return valor
        except Exception:
            print("Valor no válido. Introduzca un número entero.")


def fecha_legible(valor):
    if pd.isna(valor):
        return "Sin fecha"
    if isinstance(valor, pd.Timestamp):
        return valor.strftime("%d/%m/%Y")
    try:
        dt = pd.to_datetime(valor, errors="coerce", dayfirst=True)
        return "Sin fecha" if pd.isna(dt) else dt.strftime("%d/%m/%Y")
    except Exception:
        return str(valor)


def mostrar_sesiones_disponibles(medias_global):
    sesiones = medias_global[["Dia", "Fecha_dt", "Fecha"]].copy()

    sesiones["Fecha_sesion"] = pd.to_datetime(
        sesiones["Fecha_dt"],
        errors="coerce",
        dayfirst=True,
    )

    sin_fecha = sesiones["Fecha_sesion"].isna()
    if sin_fecha.any():
        sesiones.loc[sin_fecha, "Fecha_sesion"] = pd.to_datetime(
            sesiones.loc[sin_fecha, "Fecha"],
            errors="coerce",
            dayfirst=True,
        )

    sesiones = (
        sesiones
        .dropna(subset=["Fecha_sesion"])
        .sort_values("Fecha_sesion")
        .drop_duplicates(subset=["Fecha_sesion"])
        .reset_index(drop=True)
    )

    if sesiones.empty:
        raise ValueError("No existen sesiones con fechas válidas.")

    sesiones["Sesion"] = range(1, len(sesiones) + 1)

    print("\nSESIONES DISPONIBLES")
    print("-" * 72)
    for _, fila in sesiones.iterrows():
        print(
            f"Sesión {int(fila['Sesion']):>4} | "
            f"{fila['Fecha_sesion'].strftime('%d/%m/%Y')}"
        )
    print("-" * 72)

    return sesiones


def pedir_sesion_inicio(total_sesiones):
    return pedir_entero(
        "Seleccione el NÚMERO DE LA SESIÓN DE INICIO [1]: ",
        minimo=1,
        maximo=total_sesiones,
        defecto=1,
    )


def configurar_linea_base(medias_global):
    sesiones = mostrar_sesiones_disponibles(medias_global)
    sesion_inicio = pedir_sesion_inicio(len(sesiones))

    maximo_disponible = len(sesiones) - sesion_inicio + 1
    sesiones_solicitadas = pedir_entero(
        "¿Cuántas sesiones desea utilizar para la línea base? [7]: ",
        minimo=3,
        maximo=maximo_disponible,
        defecto=min(7, maximo_disponible),
    )

    sesion_fin = sesion_inicio + sesiones_solicitadas - 1
    fecha_inicio = sesiones.loc[
        sesiones["Sesion"] == sesion_inicio, "Fecha_sesion"
    ].iloc[0]
    fecha_fin = sesiones.loc[
        sesiones["Sesion"] == sesion_fin, "Fecha_sesion"
    ].iloc[0]

    print(
        f"\nLínea base seleccionada: sesiones "
        f"{sesion_inicio} a {sesion_fin} | "
        f"{fecha_inicio.strftime('%d/%m/%Y')} a "
        f"{fecha_fin.strftime('%d/%m/%Y')} | "
        f"{sesiones_solicitadas} sesiones.\n"
    )

    return sesiones_solicitadas, sesion_inicio, sesion_fin


def seleccionar_linea_base(medias, sesion_inicio, sesion_fin):
    validos = (
        medias
        .dropna(subset=["PAS_MEDIA", "PAD_MEDIA", "FC_MEDIA"])
        .copy()
    )

    validos["Fecha_dt"] = pd.to_datetime(
        validos["Fecha_dt"],
        errors="coerce",
        dayfirst=True,
    )

    validos = (
        validos
        .dropna(subset=["Fecha_dt"])
        .sort_values("Fecha_dt")
        .reset_index(drop=True)
    )

    validos["Sesion"] = range(1, len(validos) + 1)

    return validos[
        validos["Sesion"].between(sesion_inicio, sesion_fin)
    ].copy()


def regresion_lineal(x, y):
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    mask = np.isfinite(x) & np.isfinite(y)
    x, y = x[mask], y[mask]
    if len(x) < 2:
        return np.nan, np.nan, np.nan
    m, b = np.polyfit(x, y, 1)
    y_pred = m * x + b
    ss_res = np.sum((y - y_pred) ** 2)
    ss_tot = np.sum((y - np.mean(y)) ** 2)
    r2 = 1 - (ss_res / ss_tot) if ss_tot > 0 else np.nan
    return m, b, r2


def clasificacion_esh_home(pas, pad):
    if pd.isna(pas) or pd.isna(pad):
        return "No clasificable"
    if pas >= UMBRAL_URGENTE_PAS or pad >= UMBRAL_URGENTE_PAD:
        return "Valores muy elevados"
    if pas >= UMBRAL_HOME_PAS or pad >= UMBRAL_HOME_PAD:
        return "Presión arterial domiciliaria elevada"
    return "Presión arterial domiciliaria no elevada"


def interpretar_pas(pas):
    if pd.isna(pas):
        return "PAS no evaluable."
    if pas >= UMBRAL_URGENTE_PAS:
        return "PAS muy elevada: confirmar la medición y valorar atención sanitaria si persiste o hay síntomas."
    if pas >= UMBRAL_HOME_PAS:
        return "PAS elevada para medición domiciliaria según el umbral ESH 2023 (≥135 mmHg)."
    return "PAS dentro del rango domiciliario no hipertensivo según ESH 2023 (<135 mmHg)."


def interpretar_pad(pad):
    if pd.isna(pad):
        return "PAD no evaluable."
    if pad >= UMBRAL_URGENTE_PAD:
        return "PAD muy elevada: confirmar la medición y valorar atención sanitaria si persiste o hay síntomas."
    if pad >= UMBRAL_HOME_PAD:
        return "PAD elevada para medición domiciliaria según el umbral ESH 2023 (≥85 mmHg)."
    return "PAD dentro del rango domiciliario no hipertensivo según ESH 2023 (<85 mmHg)."


def resumen_esh_claro(pas, pad):
    if pd.isna(pas) or pd.isna(pad):
        return "No se puede clasificar la presión arterial por falta de datos válidos."
    partes = []
    partes.append("PAS elevada" if pas >= UMBRAL_HOME_PAS else "PAS no elevada")
    partes.append("PAD elevada" if pad >= UMBRAL_HOME_PAD else "PAD no elevada")
    if pas >= UMBRAL_URGENTE_PAS or pad >= UMBRAL_URGENTE_PAD:
        return "Valores muy elevados: repetir tras reposo y contactar con un profesional sanitario si se confirma o hay síntomas."
    if pas >= UMBRAL_HOME_PAS or pad >= UMBRAL_HOME_PAD:
        causa = []
        if pas >= UMBRAL_HOME_PAS:
            causa.append(f"PAS {pas:.1f} ≥ 135")
        if pad >= UMBRAL_HOME_PAD:
            causa.append(f"PAD {pad:.1f} ≥ 85")
        return "Presión arterial domiciliaria elevada por " + " y ".join(causa) + "."
    return "Presión arterial domiciliaria por debajo del punto de corte ESH 2023 de 135/85 mmHg."


def recomendacion_pa(pas, pad):
    if pd.isna(pas) or pd.isna(pad):
        return "Revisar calidad de los registros y repetir mediciones válidas."
    if pas >= UMBRAL_URGENTE_PAS or pad >= UMBRAL_URGENTE_PAD:
        return "Repetir tras 5 minutos de reposo; si persiste o hay síntomas, contactar con atención sanitaria."
    if pas >= UMBRAL_HOME_PAS or pad >= UMBRAL_HOME_PAD:
        return "Mantener registro domiciliario, revisar técnica y comentar los resultados con el profesional sanitario si la elevación persiste."
    return "Mantener seguimiento periódico, técnica correcta y hábitos cardiosaludables."


def interpretar_pendiente(variable, pendiente):
    if pd.isna(pendiente):
        return f"{variable}: tendencia no evaluable."
    if pendiente < -0.3:
        return f"{variable}: tendencia descendente ({pendiente:.2f} unidades/día)."
    if pendiente > 0.3:
        return f"{variable}: tendencia ascendente ({pendiente:.2f} unidades/día)."
    return f"{variable}: tendencia estable ({pendiente:.2f} unidades/día)."


def resumen_sujeto(medias_lb, dias_solicitados, dia_inicio, dia_fin):
    if medias_lb.empty:
        raise ValueError("No hay datos válidos en la ventana seleccionada.")
    sujeto = medias_lb["Sujeto"].iloc[0]
    pas, pad, fc = medias_lb["PAS_MEDIA"], medias_lb["PAD_MEDIA"], medias_lb["FC_MEDIA"]
    m_pas, b_pas, r2_pas = regresion_lineal(medias_lb["Dia"], pas)
    m_pad, b_pad, r2_pad = regresion_lineal(medias_lb["Dia"], pad)
    m_fc, b_fc, r2_fc = regresion_lineal(medias_lb["Dia"], fc)
    pas_m = float(pas.mean())
    pad_m = float(pad.mean())
    return {
        "Sujeto": sujeto,
        "Dias_solicitados": dias_solicitados,
        "Dia_inicio": dia_inicio,
        "Dia_fin": dia_fin,
        "Dias_validos": int(len(medias_lb)),
        "Mediciones_totales": int(medias_lb["N_mediciones"].sum()),
        "Fecha_inicio": fecha_legible(medias_lb["Fecha_dt"].iloc[0] if "Fecha_dt" in medias_lb else None),
        "Fecha_fin": fecha_legible(medias_lb["Fecha_dt"].iloc[-1] if "Fecha_dt" in medias_lb else None),
        "PAS_media": pas_m,
        "PAS_DE": float(pas.std(ddof=1)) if len(pas) > 1 else np.nan,
        "PAS_min": float(pas.min()),
        "PAS_max": float(pas.max()),
        "PAD_media": pad_m,
        "PAD_DE": float(pad.std(ddof=1)) if len(pad) > 1 else np.nan,
        "PAD_min": float(pad.min()),
        "PAD_max": float(pad.max()),
        "FC_media": float(fc.mean()),
        "FC_DE": float(fc.std(ddof=1)) if len(fc) > 1 else np.nan,
        "FC_min": float(fc.min()),
        "FC_max": float(fc.max()),
        "HRV_media": float(medias_lb["HRV_MEDIA"].mean()) if "HRV_MEDIA" in medias_lb else np.nan,
        "Presion_pulso_media": float(medias_lb["PRESION_PULSO"].mean()),
        "PAM_media": float(medias_lb["PAM"].mean()),
        "Dif_PAS_1_2_media": float(medias_lb["DIF_PAS_1_2"].mean()),
        "Dif_PAD_1_2_media": float(medias_lb["DIF_PAD_1_2"].mean()),
        "Dif_FC_1_2_media": float(medias_lb["DIF_FC_1_2"].mean()),
        "Pendiente_PAS": m_pas,
        "R2_PAS": r2_pas,
        "Pendiente_PAD": m_pad,
        "R2_PAD": r2_pad,
        "Pendiente_FC": m_fc,
        "R2_FC": r2_fc,
        "Clasificacion_ESH": clasificacion_esh_home(pas_m, pad_m),
        "Resumen_ESH": resumen_esh_claro(pas_m, pad_m),
        "Validez": "Válido" if len(medias_lb) >= dias_solicitados else f"Precaución: {len(medias_lb)} de {dias_solicitados} días válidos",
    }


def nombre_seguro(texto):
    return re.sub(r"[^A-Za-z0-9_áéíóúÁÉÍÓÚñÑ -]", "", str(texto)).replace(" ", "_")


def color_estado_pa(sigla, valor):
    if sigla == "PAS":
        if valor >= UMBRAL_URGENTE_PAS:
            return "#C00000"
        if valor >= UMBRAL_HOME_PAS:
            return "#F4B183"
        return "#A9D18E"
    if sigla == "PAD":
        if valor >= UMBRAL_URGENTE_PAD:
            return "#C00000"
        if valor >= UMBRAL_HOME_PAD:
            return "#F4B183"
        return "#A9D18E"
    return "#D9EAF7"


def crear_dashboard(resumen, carpeta):
    carpeta.mkdir(parents=True, exist_ok=True)
    sujeto = resumen["Sujeto"]
    fig = plt.figure(figsize=(13.2, 7.2))
    fig.patch.set_facecolor("white")
    fig.suptitle(f"Dashboard basal clínico | {sujeto}", fontsize=18, fontweight="bold", y=0.97)
    fig.text(0.5, 0.925, "Presión arterial domiciliaria · Orientación ESH 2023", ha="center", fontsize=11)

    datos = [
        ("PAS", resumen['PAS_media'], "mmHg", "Sistólica media"),
        ("PAD", resumen['PAD_media'], "mmHg", "Diastólica media"),
        ("FC", resumen['FC_media'], "lpm", "Frecuencia cardiaca"),
        ("PAM", resumen['PAM_media'], "mmHg", "Presión arterial media"),
        ("PP", resumen['Presion_pulso_media'], "mmHg", "Presión de pulso"),
    ]

    for i, (sigla, valor, unidad, subtitulo) in enumerate(datos):
        ax = fig.add_axes([0.045 + i * 0.19, 0.62, 0.17, 0.22])
        ax.axis("off")
        face = color_estado_pa(sigla, valor)
        rect = plt.Rectangle((0, 0), 1, 1, facecolor=face, edgecolor="#404040", linewidth=1.4, alpha=0.85)
        ax.add_patch(rect)
        ax.text(0.05, 0.80, sigla, fontsize=13, fontweight="bold")
        ax.text(0.05, 0.43, f"{valor:.1f}", fontsize=24, fontweight="bold")
        ax.text(0.68, 0.47, unidad, fontsize=10)
        ax.text(0.05, 0.16, subtitulo, fontsize=9)

    ax_txt = fig.add_axes([0.05, 0.34, 0.90, 0.20])
    ax_txt.axis("off")
    ax_txt.add_patch(plt.Rectangle((0, 0), 1, 1, fill=False, edgecolor="#808080", linewidth=1.1))
    texto = (
        f"Resultado ESH 2023: {resumen['Resumen_ESH']}\n"
        f"Ventana analizada: día {resumen['Dia_inicio']} a día {resumen['Dia_fin']} | "
        f"{resumen['Dias_validos']} días válidos | {resumen['Mediciones_totales']} mediciones.\n"
        f"Tendencia PAS: {resumen['Pendiente_PAS']:.2f} mmHg/día (R²={resumen['R2_PAS']:.2f}) · "
        f"Tendencia PAD: {resumen['Pendiente_PAD']:.2f} mmHg/día (R²={resumen['R2_PAD']:.2f}) · "
        f"Tendencia FC: {resumen['Pendiente_FC']:.2f} lpm/día (R²={resumen['R2_FC']:.2f})."
    )
    ax_txt.text(0.02, 0.82, texto, fontsize=11, va="top", linespacing=1.5)

    ax_rec = fig.add_axes([0.05, 0.12, 0.90, 0.15])
    ax_rec.axis("off")
    ax_rec.add_patch(plt.Rectangle((0, 0), 1, 1, facecolor="#F7F7F7", edgecolor="#808080", linewidth=1.0))
    ax_rec.text(0.02, 0.70, "Recomendación práctica", fontsize=11, fontweight="bold")
    ax_rec.text(0.02, 0.40, recomendacion_pa(resumen['PAS_media'], resumen['PAD_media']), fontsize=10, va="top")
    ax_rec.text(0.02, 0.13, "Nota: informe orientativo; no sustituye la valoración de un profesional sanitario.", fontsize=9, style="italic")

    ruta = carpeta / f"Dashboard_Clinico_{nombre_seguro(sujeto)}.png"
    fig.savefig(ruta, dpi=240, bbox_inches="tight")
    plt.close(fig)
    return ruta


def anotar_puntos(ax, x, y, unidad=""):
    for xi, yi in zip(x, y):
        if pd.isna(yi):
            continue
        ax.annotate(f"{yi:.1f}", (xi, yi), textcoords="offset points", xytext=(0, 8), ha="center", fontsize=8)


def crear_grafico_pa(medias, sujeto, carpeta):
    carpeta.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(11.5, 6.0))
    x = medias["Dia"].astype(float)
    ecuaciones = []
    for col, sigla in [("PAS_MEDIA", "PAS"), ("PAD_MEDIA", "PAD")]:
        y = medias[col].astype(float)
        ax.plot(x, y, marker="o", linewidth=2.4, markersize=6, label=f"{sigla} media diaria")
        anotar_puntos(ax, x, y)
        m, b, r2 = regresion_lineal(x, y)
        if not pd.isna(m):
            ax.plot(x, m*x+b, linestyle="--", linewidth=1.8, label=f"Tendencia {sigla}")
            ecuaciones.append(f"{sigla}: y={m:.2f}x+{b:.1f}; R²={r2:.2f}")
    ax.axhline(UMBRAL_HOME_PAS, linestyle=":", linewidth=1.3, label="Umbral PAS domiciliaria 135")
    ax.axhline(UMBRAL_HOME_PAD, linestyle=":", linewidth=1.3, label="Umbral PAD domiciliaria 85")
    ax.set_title(f"Evolución basal de presión arterial\n{sujeto}", fontsize=15, fontweight="bold")
    ax.set_xlabel("Día de registro basal")
    ax.set_ylabel("mmHg")
    ax.set_xticks(sorted(medias["Dia"].unique()))
    ax.grid(True, linestyle="--", alpha=0.25)
    ax.legend(fontsize=8, loc="best")
    ax.text(0.02, 0.04, "\n".join(ecuaciones), transform=ax.transAxes, fontsize=9,
            bbox=dict(boxstyle="round,pad=0.35", facecolor="white", edgecolor="gray", alpha=0.9))
    fig.tight_layout()
    ruta = carpeta / f"Grafico_PA_Profesional_{nombre_seguro(sujeto)}.png"
    fig.savefig(ruta, dpi=240, bbox_inches="tight")
    plt.close(fig)
    return ruta


def crear_grafico_fc(medias, sujeto, carpeta):
    carpeta.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(11.5, 5.2))
    x = medias["Dia"].astype(float)
    y = medias["FC_MEDIA"].astype(float)
    ax.plot(x, y, marker="o", linewidth=2.6, markersize=6, label="FC media diaria")
    anotar_puntos(ax, x, y)
    m, b, r2 = regresion_lineal(x, y)
    if not pd.isna(m):
        ax.plot(x, m*x+b, linestyle="--", linewidth=1.8, label="Tendencia FC")
        texto = f"FC: y={m:.2f}x+{b:.1f}; R²={r2:.2f}"
        ax.text(0.02, 0.05, texto, transform=ax.transAxes, fontsize=9,
                bbox=dict(boxstyle="round,pad=0.35", facecolor="white", edgecolor="gray", alpha=0.9))
    ax.set_title(f"Evolución basal de frecuencia cardiaca\n{sujeto}", fontsize=15, fontweight="bold")
    ax.set_xlabel("Día de registro basal")
    ax.set_ylabel("lpm")
    ax.set_xticks(sorted(medias["Dia"].unique()))
    ax.grid(True, linestyle="--", alpha=0.25)
    ax.legend(fontsize=9, loc="best")
    fig.tight_layout()
    ruta = carpeta / f"Grafico_FC_Profesional_{nombre_seguro(sujeto)}.png"
    fig.savefig(ruta, dpi=240, bbox_inches="tight")
    plt.close(fig)
    return ruta


def construir_tabla_vertical(row):
    pas, pad = row["PAS_media"], row["PAD_media"]
    return pd.DataFrame([
        ["PAS", "Presión arterial sistólica", f"{pas:.1f} mmHg", interpretar_pas(pas), "Seguir evolución; si ≥135 mmHg de forma persistente, comentar con profesional sanitario."],
        ["PAD", "Presión arterial diastólica", f"{pad:.1f} mmHg", interpretar_pad(pad), "Seguir evolución; si ≥85 mmHg de forma persistente, comentar con profesional sanitario."],
        ["ESH", "Clasificación domiciliaria", row["Clasificacion_ESH"], row["Resumen_ESH"], recomendacion_pa(pas, pad)],
        ["FC", "Frecuencia cardiaca", f"{row['FC_media']:.1f} lpm", interpretar_pendiente("FC", row["Pendiente_FC"]), "Interpretar junto con reposo, sueño, estrés, medicación y carga previa."],
        ["PAM", "Presión arterial media", f"{row['PAM_media']:.1f} mmHg", "Indicador hemodinámico complementario calculado como PAD + (PAS-PAD)/3.", "Interpretar siempre junto con PAS y PAD."],
        ["PP", "Presión de pulso", f"{row['Presion_pulso_media']:.1f} mmHg", "Diferencia entre PAS y PAD; aporta información complementaria.", "Vigilar cambios marcados entre días."],
        ["ΔPAS 1-2", "Diferencia PAS entre medición 1 y 2", f"{row['Dif_PAS_1_2_media']:.1f} mmHg", "Evalúa estabilidad entre tomas de la misma mañana.", "Si la diferencia es alta, reforzar reposo previo y técnica."],
        ["ΔPAD 1-2", "Diferencia PAD entre medición 1 y 2", f"{row['Dif_PAD_1_2_media']:.1f} mmHg", "Evalúa estabilidad entre tomas de la misma mañana.", "Mantener dos tomas separadas por ~1 minuto y usar la media diaria."],
        ["Tend PAS", "Pendiente de PAS", f"{row['Pendiente_PAS']:.2f} mmHg/día" if not pd.isna(row['Pendiente_PAS']) else "No evaluable", interpretar_pendiente("PAS", row["Pendiente_PAS"]), "Revisar si existe tendencia ascendente mantenida."],
        ["Tend PAD", "Pendiente de PAD", f"{row['Pendiente_PAD']:.2f} mmHg/día" if not pd.isna(row['Pendiente_PAD']) else "No evaluable", interpretar_pendiente("PAD", row["Pendiente_PAD"]), "Revisar si existe tendencia ascendente mantenida."],
        ["Validez", "Calidad de línea base", row["Validez"], "Indica si se alcanzan los días solicitados.", "Ampliar registros si no se completan los días válidos previstos."],
    ], columns=["Sigla", "Nombre completo", "Valor", "Interpretación", "Recomendaciones"])


def add_df_table(doc, df, font_size=8):
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                val = round(val, 2)
            cells[j].text = "" if pd.isna(val) else str(val)
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(font_size)
    return table


def generar_word(
    resumen_df,
    tablas_verticales,
    imagenes,
    carpeta_sujeto,
    sujeto,
):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME BASAL DE PRESIÓN ARTERIAL")
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "Orientación basada en medición domiciliaria y criterios ESH 2023"
    ).italic = True

    doc.add_paragraph(
        f"Sujeto: {sujeto}\n"
        f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    doc.add_heading("1. Criterio interpretativo", level=1)
    doc.add_paragraph(
        "El análisis utiliza la media diaria de las mediciones disponibles "
        "y resume la ventana basal seleccionada. Como orientación clínica, "
        "las guías ESH 2023 consideran elevada la presión arterial "
        "domiciliaria cuando la media es ≥135 mmHg para PAS y/o ≥85 mmHg "
        "para PAD. Este informe no sustituye la valoración médica."
    )

    doc.add_heading("2. Resumen general", level=1)
    cols = [
        "Sujeto",
        "Dias_solicitados",
        "Dia_inicio",
        "Dia_fin",
        "Dias_validos",
        "Mediciones_totales",
        "Fecha_inicio",
        "Fecha_fin",
        "PAS_media",
        "PAD_media",
        "FC_media",
        "PAM_media",
        "Clasificacion_ESH",
        "Validez",
    ]

    tabla_resumen = resumen_df[cols].copy()
    tabla_resumen.columns = [
        "Sujeto",
        "Sesiones solicitadas",
        "Sesión inicio",
        "Sesión fin",
        "Sesiones válidas",
        "Mediciones",
        "Fecha inicio",
        "Fecha fin",
        "PAS",
        "PAD",
        "FC",
        "PAM",
        "Clasificación ESH",
        "Validez",
    ]
    add_df_table(doc, tabla_resumen, font_size=7)

    doc.add_heading("3. Análisis individual", level=1)
    row = resumen_df.iloc[0]

    if sujeto in imagenes and "dashboard" in imagenes[sujeto]:
        doc.add_picture(
            str(imagenes[sujeto]["dashboard"]),
            width=Inches(9.6),
        )

    doc.add_paragraph("Tabla interpretativa vertical")
    add_df_table(doc, tablas_verticales[sujeto], font_size=7)

    doc.add_paragraph("Evolución de presión arterial")
    if sujeto in imagenes and "pa" in imagenes[sujeto]:
        doc.add_picture(
            str(imagenes[sujeto]["pa"]),
            width=Inches(8.8),
        )

    doc.add_paragraph("Evolución de frecuencia cardiaca")
    if sujeto in imagenes and "fc" in imagenes[sujeto]:
        doc.add_picture(
            str(imagenes[sujeto]["fc"]),
            width=Inches(8.8),
        )

    doc.add_heading("4. Recomendaciones de técnica de medición", level=1)
    recomendaciones = [
        "Realizar las mediciones por la mañana, antes del desayuno y antes de actividad física intensa.",
        "Permanecer sentado al menos 5 minutos antes de la primera toma.",
        "Usar manguito adecuado, brazo apoyado a la altura del corazón y evitar hablar durante la medición.",
        "Registrar dos tomas separadas aproximadamente por un minuto y analizar la media diaria.",
        "Consultar con un profesional sanitario ante medias domiciliarias persistentemente elevadas o valores muy elevados.",
    ]

    for recomendacion in recomendaciones:
        doc.add_paragraph(recomendacion)

    doc.add_heading("5. Referencias", level=1)
    referencias = [
        "European Society of Hypertension. 2023 ESH Guidelines for the management of arterial hypertension. Journal of Hypertension.",
        "Stergiou, G. S., et al. Home blood pressure monitoring: methodology, clinical relevance and practical application. Journal of Hypertension.",
        "European Society of Cardiology. 2024 ESC Guidelines for the management of elevated blood pressure and hypertension. European Heart Journal.",
    ]

    for referencia in referencias:
        doc.add_paragraph(referencia)

    ruta = carpeta_sujeto / (
        f"Informe_Basal_Presion_Arterial_{nombre_seguro(sujeto)}.docx"
    )
    doc.save(ruta)
    return ruta


def guardar_excel(
    resumen_df,
    medias_df,
    raw_df,
    tablas_verticales,
    errores,
    sesiones_linea_base,
    sesion_inicio,
    sesion_fin,
    carpeta_sujeto,
    sujeto,
):
    ruta = carpeta_sujeto / (
        f"Resultados_Basal_Presion_Arterial_{nombre_seguro(sujeto)}.xlsx"
    )

    with pd.ExcelWriter(ruta, engine="openpyxl") as writer:
        resumen_df.to_excel(
            writer,
            sheet_name="01_RESUMEN",
            index=False,
        )
        medias_df.to_excel(
            writer,
            sheet_name="02_MEDIAS_DIARIAS",
            index=False,
        )
        raw_df.to_excel(
            writer,
            sheet_name="03_DATOS_ORIGINALES",
            index=False,
        )

        tablas_verticales[sujeto].to_excel(
            writer,
            sheet_name="04_INTERPRETACION",
            index=False,
        )

        pd.DataFrame(
            [
                ["Sesiones solicitadas", sesiones_linea_base],
                ["Sesión inicio", sesion_inicio],
                ["Sesión fin", sesion_fin],
                [
                    "Criterio ESH 2023",
                    "PA domiciliaria elevada si media PAS >=135 y/o PAD >=85 mmHg",
                ],
                [
                    "Advertencia",
                    "Informe orientativo; no sustituye valoración médica",
                ],
            ],
            columns=["Parámetro", "Valor"],
        ).to_excel(
            writer,
            sheet_name="05_CONFIGURACION",
            index=False,
        )

        pd.DataFrame(
            [
                ["PAS", "Presión arterial sistólica", "mmHg"],
                ["PAD", "Presión arterial diastólica", "mmHg"],
                ["FC", "Frecuencia cardiaca", "lpm"],
                [
                    "PAM",
                    "Presión arterial media = PAD + (PAS-PAD)/3",
                    "mmHg",
                ],
                ["PP", "Presión de pulso = PAS - PAD", "mmHg"],
                [
                    "Pendiente",
                    "Cambio estimado por sesión mediante regresión lineal",
                    "unidad/sesión",
                ],
                [
                    "R²",
                    "Proporción de varianza explicada por la tendencia lineal",
                    "0-1",
                ],
            ],
            columns=["Variable", "Definición", "Unidad"],
        ).to_excel(
            writer,
            sheet_name="06_DICCIONARIO",
            index=False,
        )

        if errores:
            pd.DataFrame(errores).to_excel(
                writer,
                sheet_name="07_ERRORES",
                index=False,
            )

    return ruta



def configurar_consola():
    for nombre in ("stdin", "stdout", "stderr"):
        flujo = getattr(sys, nombre, None)
        if hasattr(flujo, "reconfigure"):
            try:
                flujo.reconfigure(encoding="utf-8")
            except Exception:
                pass


def listar_archivos_entrada():
    extensiones = {".xlsx", ".xls", ".csv"}
    return sorted(
        [
            ruta
            for ruta in CARPETA_DATOS.iterdir()
            if ruta.is_file() and ruta.suffix.lower() in extensiones
        ],
        key=lambda ruta: ruta.name.lower(),
    )


def leer_csv_robusto(ruta_csv):
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        for separador in (None, ";", ",", "\t"):
            try:
                opciones = {"encoding": encoding}
                if separador is None:
                    opciones.update({"sep": None, "engine": "python"})
                else:
                    opciones["sep"] = separador

                df = pd.read_csv(ruta_csv, **opciones)
                if len(df.columns) >= 2:
                    return df
            except Exception:
                continue

    raise ValueError(f"No se pudo leer el archivo CSV: {ruta_csv.name}")


def convertir_csv_a_excel(ruta_csv):
    CARPETA_CONVERTIDOS.mkdir(parents=True, exist_ok=True)
    destino = CARPETA_CONVERTIDOS / f"{ruta_csv.stem}.xlsx"
    leer_csv_robusto(ruta_csv).to_excel(destino, index=False)
    print(f"CSV convertido: {ruta_csv.name} -> {destino}")
    return destino


def preparar_archivo(ruta):
    if ruta.suffix.lower() == ".csv":
        return convertir_csv_a_excel(ruta)
    return ruta


def cargar_entradas(archivos):
    entradas = []
    errores = []

    print("\nAnalizando archivos...")

    for ruta_original in archivos:
        try:
            ruta_lectura = preparar_archivo(ruta_original)
            raw = leer_excel(ruta_lectura)
            sujeto = str(raw["Sujeto"].iloc[0]).strip()

            entradas.append(
                {
                    "Sujeto": sujeto,
                    "Archivo": ruta_original.name,
                    "Formato": ruta_original.suffix.upper().lstrip("."),
                    "Raw": raw,
                    "Ruta_original": ruta_original,
                    "Ruta_lectura": ruta_lectura,
                }
            )
        except Exception as exc:
            errores.append(
                {
                    "Archivo": ruta_original.name,
                    "Error": str(exc),
                }
            )
            print(f"ERROR leyendo {ruta_original.name}: {exc}")

    if not entradas:
        raise RuntimeError("No se pudo cargar ningún archivo válido.")

    return entradas, errores


def construir_resumen_entradas(entradas):
    resumen = []

    for entrada in entradas:
        raw = entrada["Raw"]
        sesiones = (
            pd.to_datetime(
                raw["Fecha_dt"],
                errors="coerce",
                dayfirst=True,
            )
            .dropna()
            .dt.normalize()
            .nunique()
        )

        resumen.append(
            {
                "Sujeto": entrada["Sujeto"],
                "Archivo": entrada["Archivo"],
                "Formato": entrada["Formato"],
                "Registros": len(raw),
                "Sesiones": int(sesiones),
                "Raw": raw,
            }
        )

    return resumen


def mostrar_entradas(resumen_entradas):
    print("\n" + "=" * 78)
    print("ARCHIVOS / SUJETOS DISPONIBLES")
    print("=" * 78)

    for indice, entrada in enumerate(resumen_entradas, start=1):
        print(f"{indice:>3}. {entrada['Sujeto']}")
        print(
            f"     Archivo  : {entrada['Archivo']}\n"
            f"     Registros: {entrada['Registros']} | "
            f"Sesiones: {entrada['Sesiones']} | "
            f"Formato: {entrada['Formato']}"
        )

    print("-" * 78)
    print("Seleccione uno: 1")
    print("Seleccione varios: 1,2")
    print("Seleccione todos: A")
    print("Salir: 0")


def seleccionar_entradas(resumen_entradas):
    while True:
        mostrar_entradas(resumen_entradas)
        entrada = input("\nSelección: ").strip()

        if entrada == "0":
            return []

        if entrada.lower() == "a":
            return resumen_entradas

        try:
            indices = []
            for parte in entrada.split(","):
                indice = int(parte.strip())
                if indice < 1 or indice > len(resumen_entradas):
                    raise ValueError
                if indice not in indices:
                    indices.append(indice)

            if not indices:
                raise ValueError

            return [resumen_entradas[indice - 1] for indice in indices]

        except Exception:
            print("Selección no válida. Use 1, 1,2, A o 0.")


def carpeta_salida_entrada(sujeto, archivo):
    base = nombre_seguro(sujeto)
    origen = nombre_seguro(Path(archivo).stem)
    return CARPETA_RESULTADOS / f"{base}__{origen}"


def main():
    try:
        configurar_consola()

        CARPETA_DATOS.mkdir(parents=True, exist_ok=True)
        CARPETA_RESULTADOS.mkdir(parents=True, exist_ok=True)
        CARPETA_CONVERTIDOS.mkdir(parents=True, exist_ok=True)

        print("=" * 78)
        print("SportsLabResearch - Blood Pressure Analyzer v0.6.7")
        print("=" * 78)

        archivos = listar_archivos_entrada()
        if not archivos:
            raise FileNotFoundError(
                "No hay archivos .xlsx, .xls o .csv en data/input."
            )

        entradas, errores_globales = cargar_entradas(archivos)
        resumen_entradas = construir_resumen_entradas(entradas)
        seleccionados = seleccionar_entradas(resumen_entradas)

        if not seleccionados:
            print("\nProceso cancelado.")
            return

        resultados_generados = []

        for elemento in seleccionados:
            sujeto = elemento["Sujeto"]
            archivo = elemento["Archivo"]
            raw = elemento["Raw"].copy()

            print("\n" + "=" * 78)
            print(f"SUJETO: {sujeto}")
            print(f"ARCHIVO: {archivo}")
            print("=" * 78)

            try:
                carpeta_sujeto = carpeta_salida_entrada(
                    sujeto,
                    archivo,
                )
                carpeta_figuras = carpeta_sujeto / "figures"

                carpeta_sujeto.mkdir(
                    parents=True,
                    exist_ok=True,
                )
                carpeta_figuras.mkdir(
                    parents=True,
                    exist_ok=True,
                )

                medias = calcular_medias_diarias(raw)
                medias["Linea_base"] = False

                (
                    sesiones_linea_base,
                    sesion_inicio,
                    sesion_fin,
                ) = configurar_linea_base(medias)

                medias_lb = seleccionar_linea_base(
                    medias,
                    sesion_inicio,
                    sesion_fin,
                )

                if medias_lb.empty:
                    raise ValueError(
                        "No hay sesiones válidas dentro de la línea base seleccionada."
                    )

                fechas_lb = pd.to_datetime(
                    medias_lb["Fecha_dt"],
                    errors="coerce",
                    dayfirst=True,
                ).dt.normalize()

                fechas_medias = pd.to_datetime(
                    medias["Fecha_dt"],
                    errors="coerce",
                    dayfirst=True,
                ).dt.normalize()

                medias.loc[
                    fechas_medias.isin(fechas_lb),
                    "Linea_base",
                ] = True

                resumen = resumen_sujeto(
                    medias_lb,
                    sesiones_linea_base,
                    sesion_inicio,
                    sesion_fin,
                )

                resumen_df = pd.DataFrame([resumen])
                tabla_vertical = construir_tabla_vertical(resumen)
                tablas_verticales = {sujeto: tabla_vertical}

                imagenes = {
                    sujeto: {
                        "dashboard": crear_dashboard(
                            resumen,
                            carpeta_figuras,
                        ),
                        "pa": crear_grafico_pa(
                            medias_lb,
                            sujeto,
                            carpeta_figuras,
                        ),
                        "fc": crear_grafico_fc(
                            medias_lb,
                            sujeto,
                            carpeta_figuras,
                        ),
                    }
                }

                for dataframe in (raw, medias, resumen_df):
                    for columna in dataframe.select_dtypes(
                        include=[np.number]
                    ).columns:
                        dataframe[columna] = dataframe[columna].round(3)

                errores_entrada = [
                    error
                    for error in errores_globales
                    if error.get("Archivo") == archivo
                ]

                ruta_excel = guardar_excel(
                    resumen_df,
                    medias,
                    raw,
                    tablas_verticales,
                    errores_entrada,
                    sesiones_linea_base,
                    sesion_inicio,
                    sesion_fin,
                    carpeta_sujeto,
                    sujeto,
                )

                ruta_word = generar_word(
                    resumen_df,
                    tablas_verticales,
                    imagenes,
                    carpeta_sujeto,
                    sujeto,
                )

                resultados_generados.append(
                    {
                        "Sujeto": sujeto,
                        "Archivo": archivo,
                        "Carpeta": carpeta_sujeto,
                        "Excel": ruta_excel,
                        "Word": ruta_word,
                    }
                )

                print(f"Procesado correctamente: {sujeto}")
                print(f"Archivo: {archivo}")
                print(f"Carpeta: {carpeta_sujeto}")
                print(f"Excel: {ruta_excel}")
                print(f"Word: {ruta_word}")

            except Exception as exc:
                print(
                    f"ERROR procesando {sujeto} "
                    f"({archivo}): {exc}"
                )

        if not resultados_generados:
            raise RuntimeError(
                "No se pudo procesar ninguna selección."
            )

        print("\n" + "=" * 78)
        print("ANÁLISIS FINALIZADO")
        print("=" * 78)
        print(
            f"Entradas procesadas: {len(resultados_generados)}"
        )

        for resultado in resultados_generados:
            print(
                f"- {resultado['Sujeto']} | "
                f"{resultado['Archivo']} | "
                f"{resultado['Carpeta']}"
            )

    except Exception as exc:
        print("\n[ERROR]")
        print(str(exc))
        try:
            input("\nPulse Enter para cerrar...")
        except Exception:
            pass
        sys.exit(1)


if __name__ == "__main__":
    main()
